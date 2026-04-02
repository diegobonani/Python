from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def configurar_fonte(run, tamanho=11, negrito=False, cor=None, nome_fonte='Arial'):
    run.font.name = nome_fonte
    run.font.size = Pt(tamanho)
    run.bold = negrito
    if cor:
        run.font.color.rgb = cor
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), nome_fonte)
    rFonts.set(qn('w:hAnsi'), nome_fonte)
    rPr.append(rFonts)

def adicionar_borda_inferior(paragrafo):
    p = paragrafo._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def criar_cv_english():
    document = Document()

    # --- MARGINS ---
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # --- HEADER ---
    p_nome = document.add_paragraph()
    p_nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_nome = p_nome.add_run('DIEGO BONANI ARANTES')
    configurar_fonte(run_nome, tamanho=22, negrito=True)

    p_info = document.add_paragraph()
    p_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_info.paragraph_format.space_after = Pt(12)
    
    def add_contact(text, is_link=False):
        r = p_info.add_run(text)
        color = RGBColor(0, 51, 153) if is_link else RGBColor(0, 0, 0)
        configurar_fonte(r, tamanho=10, cor=color)

    add_contact('Edmonton, AB (Intended Destination)  |  +55 (12) 98851-1752\n')
    add_contact('diego.bonani20@gmail.com', True)
    add_contact('  |  ')
    add_contact('linkedin.com/in/diegobonani', True)
    add_contact('  |  ')
    add_contact('github.com/diegobonani', True)

    # --- SECTION TITLE FUNCTION ---
    def add_section(title):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title.upper())
        configurar_fonte(r, tamanho=11, negrito=True)
        adicionar_borda_inferior(p)

    # --- PROFESSIONAL SUMMARY ---
    add_section('Professional Summary')
    p_sum = document.add_paragraph()
    p_sum.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    text_sum = (
        "Results-oriented IT Infrastructure Analyst with over 6 years of experience in retail environments. "
        "Specialized in managing high-availability networks (Mikrotik/Linux), troubleshooting ERP systems (TOTVS Protheus), "
        "and automating legacy workflows. Recently developed a Python/Bash monitoring solution that reduced POS diagnosis "
        "time by 90% across 17 branches. Skilled in bridging the gap between hardware support and software development "
        "(JSON/API integration) to drive operational efficiency."
    )
    configurar_fonte(p_sum.add_run(text_sum), tamanho=10.5)

    # --- TECHNICAL SKILLS ---
    add_section('Technical Skills')
    skills = [
        ("Infrastructure & Networking:", "Mikrotik RouterOS (VPN, Firewall, Queues), Linux (Debian/RedHat), Windows Server, SSH Tunneling, WSL."),
        ("Automation & Scripting:", "Python (Pandas, Requests), Bash Scripting, Task Scheduling, JSON/Excel Data Extraction."),
        ("Retail Systems:", "TOTVS Protheus (Consinco Line), POS Hardware & Peripherals, SAT Fiscal Modules, ERP Troubleshooting."),
        ("Tools & Support:", "GLPI (Admin & Customization), SQL (Basic Queries), Git/GitHub, Office 365 Administration.")
    ]
    for cat, desc in skills:
        p = document.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r_cat = p.add_run(cat + " ")
        configurar_fonte(r_cat, tamanho=10.5, negrito=True)
        r_desc = p.add_run(desc)
        configurar_fonte(r_desc, tamanho=10.5)

    # --- PROFESSIONAL EXPERIENCE ---
    add_section('Professional Experience')

    def add_job(role, company, location, date, bullets):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(0)
        
        # Line 1: Role and Date
        r_role = p.add_run(role)
        configurar_fonte(r_role, tamanho=11, negrito=True)
        r_sep = p.add_run(f"\t{date}")
        configurar_fonte(r_sep, tamanho=11, negrito=False)
        p.tab_stops.add_tab_stop(Inches(7), 2) # Right align date

        # Line 2: Company
        p2 = document.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        r_comp = p2.add_run(f"{company} | {location}")
        configurar_fonte(r_comp, tamanho=10.5, negrito=True, cor=RGBColor(80, 80, 80))

        # Bullets
        for b in bullets:
            pb = document.add_paragraph(style='List Bullet')
            pb.paragraph_format.space_after = Pt(1)
            rb = pb.add_run(b)
            configurar_fonte(rb, tamanho=10.5)

    # Job 1 - VillaSimpatia
    add_job(
        "IT Infrastructure Analyst", 
        "Comercial VillaSimpatia (Retail Chain)", 
        "Brazil", 
        "Aug 2024 – Present",
        [
            "Manage IT infrastructure for 17 supermarket branches, ensuring high availability of POS terminals and Mikrotik network appliances.",
            "Key Achievement: Developed a custom CLI automation tool using Python/Bash/WSL to bulk-monitor POS health (CPU/RAM/Ping) and export status to JSON/Excel. Reduced daily checkup time from 1 hour to 2 minutes.",
            "Optimized the GLPI ticketing system by implementing conditional frontend logic, reducing user error in toner requests by 40%.",
            "Provide Tier 3 support for TOTVS Protheus ERP (Consinco), troubleshooting database connectivity and fiscal modules."
        ]
    )

    # Job 2 - Freelance
    add_job(
        "Full Stack Web Developer (Freelance)", 
        "Self-Employed", 
        "Remote", 
        "Aug 2023 – Jul 2024",
        [
            "Developed web applications using React.js, Next.js, and Node.js for various clients.",
            "Maintained and refactored legacy PHP/jQuery codebases, improving performance and security.",
            "Designed relational database schemas and integrated RESTful APIs."
        ]
    )

    # Job 3 - Semar
    add_job(
        "IT Support Technician", 
        "Semar Supermarkets", 
        "Brazil", 
        "Feb 2021 – Aug 2023",
        [
            "Provided on-site and remote support for 9 stores across the state, managing hardware maintenance and network troubleshooting.",
            "Administered Bluesoft ERP users, stock movements, and fiscal printers (SAT).",
            "Responsible for server backups and preventive maintenance of POS hardware."
        ]
    )

    # Job 4 - Internship
    add_job(
        "IT Intern", 
        "Regional Attorney's Office (PGE)", 
        "Brazil", 
        "Aug 2018 – Jan 2020",
        [
            "Served as the sole IT point of contact for the unit, demonstrating autonomy and self-management.",
            "Managed structured cabling, hardware repairs, and provided support for the SOFTPLAN legal system."
        ]
    )

    # --- EDUCATION ---
    add_section('Education')
    p_edu = document.add_paragraph()
    r_deg = p_edu.add_run("Bachelor of Computer Science\n")
    configurar_fonte(r_deg, tamanho=11, negrito=True)
    r_uni = p_edu.add_run("Anhanguera Educacional | Concluded: June 2020")
    configurar_fonte(r_uni, tamanho=10.5)

    # --- CERTIFICATIONS ---
    add_section('Certifications & Languages')
    certs = [
        "LPI Linux Essentials – Dio.me",
        "Python (Basic & Intermediate) – Fundação Bradesco",
        "Information Security – Fundação Bradesco",
        "Languages: Portuguese (Native), English (Advanced - Ready for work environments)"
    ]
    for c in certs:
        p = document.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(0)
        configurar_fonte(p.add_run(c), tamanho=10.5)

    document.save('Diego_Bonani_CV_English.docx')
    print("CV em Inglês gerado com sucesso!")

if __name__ == '__main__':
    criar_cv_english()